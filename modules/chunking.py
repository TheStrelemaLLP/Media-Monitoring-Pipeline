import os
import sys
import io
import re
import logging
from docx import Document
import docx

# Add project root to path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

# ---------------------------
# Logging Configuration
# ---------------------------
LOGS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
os.makedirs(LOGS_DIR, exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),  # Console logs
        logging.FileHandler(os.path.join(LOGS_DIR, "chunking.log"), mode="a", encoding="utf-8")  # File logs
    ]
)
logger = logging.getLogger("Chunking")


class Chunking:
    def get_docx_size_kb(self, doc):
        """Return the in-memory size of the Document in KB."""
        buffer = io.BytesIO()
        doc.save(buffer)
        size = len(buffer.getvalue()) / 1024
        logger.debug(f"Calculated DOCX size: {size:.2f} KB")
        return size

    def copy_paragraph(self, src_para, dest_doc):
        """Copy paragraph from source to destination doc, preserving formatting and embedded hyperlinks."""
        new_para = dest_doc.add_paragraph()
        new_para.style = src_para.style

        # Copy all runs (bold, italic, underline, font)
        for run in src_para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

        # Manually check and add hyperlinks if present in XML
        hyperlinks = src_para._element.xpath('.//w:hyperlink')
        for hlink in hyperlinks:
            r_id = hlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if r_id:
                rel = src_para.part.rels[r_id]
                new_para.add_run(f" {rel.target_ref}")  # Add the URL at the end
                logger.debug(f"Copied hyperlink: {rel.target_ref}")

    def is_article_start(self, paragraph):
        """Detect start of a new article (case insensitive)."""
        result = bool(re.match(r"(?i)^article\s+\d+", paragraph.text.strip()))
        if result:
            logger.debug(f"Detected article start: '{paragraph.text.strip()}'")
        return result

    def split_docx_by_article_chunks(self, input_path, chunk_size_kb, output_dir="output", leader=""):
        """Split DOCX into chunks at article boundaries, each under a size limit, named with leader."""
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Loading document: {input_path}")

        original_doc = Document(input_path)
        paragraphs = [p for p in original_doc.paragraphs if p.text.strip()]
        logger.info(f"Found {len(paragraphs)} non-empty paragraphs")

        # Group paragraphs into article blocks
        articles = []
        current_article = []

        for para in paragraphs:
            if self.is_article_start(para) and current_article:
                articles.append(current_article)
                current_article = []
            current_article.append(para)
        if current_article:
            articles.append(current_article)

        logger.info(f"Grouped into {len(articles)} article(s)")

        chunks = []
        chunk_index = 1
        current_doc = Document()
        current_doc._body.clear_content()

        for article in articles:
            # Create test doc to check size with this article added
            test_doc = Document()
            test_doc._body.clear_content()
            for p in current_doc.paragraphs:
                self.copy_paragraph(p, test_doc)
            for p in article:
                self.copy_paragraph(p, test_doc)

            size_kb = self.get_docx_size_kb(test_doc)

            if size_kb <= chunk_size_kb:
                for p in article:
                    self.copy_paragraph(p, current_doc)
            else:
                # Save current chunk
                chunk_filename = f"{leader}_chunk_{chunk_index}.docx" if leader else f"chunk_{chunk_index}.docx"
                chunk_path = os.path.join(output_dir, chunk_filename)
                current_doc.save(chunk_path)
                logger.info(f"Saved chunk: {chunk_path} ({self.get_docx_size_kb(current_doc):.2f} KB)")
                chunks.append(chunk_path)
                chunk_index += 1

                # Start new document with current article
                current_doc = Document()
                current_doc._body.clear_content()
                for p in article:
                    self.copy_paragraph(p, current_doc)

        # Save final chunk
        if current_doc.paragraphs:
            chunk_filename = f"{leader}_chunk_{chunk_index}.docx" if leader else f"chunk_{chunk_index}.docx"
            chunk_path = os.path.join(output_dir, chunk_filename)
            current_doc.save(chunk_path)
            logger.info(f"Saved chunk: {chunk_path} ({self.get_docx_size_kb(current_doc):.2f} KB)")
            chunks.append(chunk_path)

        logger.info(f"Done. {len(chunks)} chunk(s) created at '{output_dir}'.")
        return chunks
