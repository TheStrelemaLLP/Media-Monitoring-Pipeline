import os
import sys
from datetime import datetime

# Add parent directory to path for local imports
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import pandas as pd
from docx import Document
from apify_client import ApifyClient

from config import Config
from utils import setup_logger

logger = setup_logger("FacebookDocs")

class FacebookDocsError(Exception):
    """Base exception for FacebookDocs errors."""
    pass

class Facebook_doc:
    """
    A class to collect Facebook posts using Apify and save them to a DOCX file.
    """
    def __init__(self):
        """Initializes the Apify client."""
        self.client = ApifyClient(Config.apify_key)

    def collect_fb_posts(self, facebook_profile_url, start_date=None, end_date=None):
        """
        Runs the Apify actor to collect posts from a Facebook URL.
        """
        logger.info(f"Processing Facebook posts from {facebook_profile_url} | From: {start_date} To: {end_date}")

        run_input = {
            "startUrls": [{"url": facebook_profile_url}],
            "resultsType": "posts",
            "resultsLimit": 250,  # Number of posts
            "proxyConfig": {"useApifyProxy": True},
        }

        # Run the actor and wait for it to finish
        run = self.client.actor("KoJrdxJCTtpon81KY").call(run_input=run_input)

        # Fetch posts from the actor's dataset
        dataset_items = list(self.client.dataset(run["defaultDatasetId"]).iterate_items())
        logger.info(f"Fetched {len(dataset_items)} posts for {facebook_profile_url}")
        
        return dataset_items

    def filter_posts_datewise(self, df, start_date, end_date):
        """
        Filters a DataFrame based on a date range using pandas datetime.
        """
        # Ensure 'time' column exists before proceeding
        if 'time' not in df.columns:
            logger.error("DataFrame is missing the 'time' column.")
            return pd.DataFrame() # Return an empty DataFrame

        df['time'] = pd.to_datetime(df['time'], errors='coerce')
        # df.dropna(subset=['time'], inplace=True)

        start_date_obj = pd.to_datetime(start_date).date()
        end_date_obj = pd.to_datetime(end_date).date()

        filtered_df = df[
            (df['time'].dt.date >= start_date_obj) & (df['time'].dt.date <= end_date_obj)
        ]
        return filtered_df

    def save_fb_docx(self, fb_posts, output_dir, leader, start_date, end_date):
        """
        Filters posts by date and saves them to a .docx file.
        """
        if not fb_posts:
            logger.warning(f"No posts found for {leader} to save.")
            return None

        fb_posts_df = pd.DataFrame(fb_posts)
        fb_posts_df.to_excel(output_dir+"/fb_excel.xlsx", index= False)

        if fb_posts_df.empty:
            logger.info(f"No data for {leader} to process after creating DataFrame.")
            return None

        filtered_posts_df = self.filter_posts_datewise(fb_posts_df, start_date, end_date)

        if filtered_posts_df.empty:
            logger.info(f"No posts found for {leader} within the date range {start_date} to {end_date}.")
            return None

        doc = Document()
        
        article_n = 0
        for item in filtered_posts_df.to_dict('records'):
            article_n = article_n + 1
            post_time = item['time'].strftime('%Y-%m-%d %H:%M')
            heading_text = f"Facebook Post {article_n} ({post_time}):\n"
            content_text = f"Content:\n{leader} posted:\n\"{item.get('text', 'No text content.')}\""

            doc.add_paragraph(heading_text + content_text)
            doc.add_paragraph()

        # --- START OF FIX ---
        # Convert dates to datetime objects to handle them consistently
        start_date_obj = pd.to_datetime(start_date)
        end_date_obj = pd.to_datetime(end_date)
        
        # Format the dates into a clean, file-safe string (YYYY-MM-DD)
        start_date_safe_str = start_date_obj.strftime('%Y-%m-%d')
        end_date_safe_str = end_date_obj.strftime('%Y-%m-%d')

        # Use the safe strings to create the filename
        output_filename = f"{leader}_fb_output_{start_date_safe_str}_to_{end_date_safe_str}.docx"
        # --- END OF FIX ---

        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)

        doc.save(output_path)
        logger.info(f"Formatted Word document with {len(filtered_posts_df)} posts saved as {output_path}")
        return output_path
        
