import os
import sys
import pathlib
import traceback
from urllib.parse import urlparse

# Add parent directory to path for local imports
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import pandas as pd
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.common.exceptions import WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import RGBColor
from tqdm import tqdm
from requests.exceptions import RequestException

from utils import setup_logger
from modules.link_extraction import Link_extraction

logger = setup_logger("NewsScraper")

# Custom exceptions
class NewsScraperError(Exception):
    """Base exception for NewsScraper errors."""
    pass

class InvalidLeaderError(NewsScraperError):
    """Raised when an invalid leader is specified."""
    pass

class FileAccessError(NewsScraperError):
    """Raised when file operations fail."""
    pass

class WebScrapingError(NewsScraperError):
    """Raised when web scraping fails."""
    pass

# Placeholder for utility functions
def clean_text(text: str) -> str:
    """Clean text by removing extra whitespace and special characters."""
    return ' '.join(text.strip().split()) if text else ''

def ensure_heading_style(doc: Document, style_name: str, level: int):
    """Ensure heading style exists in the document."""
    try:
        return doc.styles[style_name]
    except KeyError:
        logger.error(f"Style '{style_name}' not found in document.")
        raise NewsScraperError(f"Document style '{style_name}' is missing.")

def is_valid_url(url: str) -> bool:
    """Check if a string is a valid URL."""
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except ValueError:
        return False

class NewsScraper:
    def __init__(self, file_path: str, output_folder: str, leader: str, chunk_size_kb: int, i_file_path: str, link_extraction: object):
        """Initialize NewsScraper with input file, output folder, and leader."""
        logger.info("Initializing NewsScraper.")
        self.link_extraction = link_extraction
        self.i_file_path = i_file_path
        self.file_path = file_path
        self.output_folder = output_folder
        self.leader = leader
        self.chunk_size_kb = chunk_size_kb
        self.news_links = []
        self.extracted_data = []
        self.failed_links = []
        self.driver = None
        self.docx_output, self.excel_output = self._get_output_filenames()
        self.LEADERS = self.get_leaders_variations(i_file_path, leader) 
        
        # # Validate inputs
        # if not os.path.exists(file_path):
        #     logger.error(f"Input file '{file_path}' does not exist.")
        #     raise FileAccessError(f"Input file '{file_path}' does not exist.")
        # if not self.leader in self.LEADERS:
        #     logger.error(f"Invalid leader '{leader}'. Available: {', '.join(self.LEADERS.keys())}")
        #     raise InvalidLeaderError(
        #         f"Invalid leader '{leader}'. Choose from: {', '.join(self.LEADERS.keys())}"
        #     )
        logger.debug(f"NewsScraper initialized with leader: {leader}")

    def _get_output_filenames(self) -> tuple:
        """Generate output file names based on input file and leader."""
        try:
            file_name = pathlib.Path(self.file_path).stem
            docx_output = f"{file_name}_{self.leader}_News_Content.docx"
            excel_output = f"{file_name}_{self.leader}_failed_links.xlsx"
            logger.debug(f"Output filenames set: {docx_output}, {excel_output}")
            return docx_output, excel_output
        except Exception as e:
            logger.error(f"Error generating output filenames: {str(e)}")
            raise NewsScraperError("Failed to generate output filenames.")

    def get_leaders_variations(self, file_path, sheet_name):
        """
        Extracts all non-empty values from 'English' and 'Marathi' columns in a given Excel sheet.

        Parameters:
        - file_path (str): Path to the Excel file.
        - sheet_name (str): Name of the sheet to process.

        Returns:
        - dict: {
            sheet_name: {
                "English": [list of non-empty values from English column],
                "Marathi": [list of non-empty values from Marathi column]
            }
        }

        Raises:
        - Logs and re-raises exceptions on failure.
        """
        try:
            logger.info(f"Reading Excel file: {file_path}, Sheet: {sheet_name}")
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            if 'English' not in df.columns and 'Marathi' not in df.columns:
                logger.error("Neither 'English' nor 'Marathi' columns found in the sheet.")
                raise KeyError("Missing required columns: 'English' and/or 'Marathi'.")

            english_list = df['English'].dropna().astype(str).tolist() if 'English' in df.columns else []
            marathi_list = df['Marathi'].dropna().astype(str).tolist() if 'Marathi' in df.columns else []

            LEADERS = {
                sheet_name: {
                    "English": english_list,
                    "Marathi": marathi_list
                }
            }

            logger.info(f"Extracted {len(english_list)} English and {len(marathi_list)} Marathi names from '{sheet_name}'")
            return LEADERS

        except FileNotFoundError:
            logger.exception(f"File not found: {file_path}")
            raise

        except ValueError:
            logger.exception(f"Sheet '{sheet_name}' not found in the Excel file.")
            # raise
            pass

        except KeyError as ke:
            logger.exception(f"Missing column in sheet: {ke}")
            raise

        except Exception as e:
            logger.exception("Unexpected error while extracting leader variations.")
            raise

    def _extract_static_page(self, url: str) -> tuple:
        """Extract content from a static webpage."""
        logger.debug(f"Attempting static scrape for URL: {url}")
        try:
            response = requests.get(url, timeout=1)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            headline = soup.find("h1").text if soup.find("h1") and soup.find("h1").text.strip() else \
                      soup.find("h2").text if soup.find("h2") and soup.find("h2").text.strip() else "No headline"
            paragraphs = [p.text for p in soup.find_all("p")]
            article = "\n".join(paragraphs)
            logger.debug(f"Static scrape successful for {url}")
            return headline, article
        except RequestException as e:
            logger.warning(f"Static scrape failed for {url}: {str(e)}")
            return None, None
        except Exception as e:
            logger.error(f"Unexpected error in static scrape for {url}: {str(e)}\n{traceback.format_exc()}")
            return None, None    

    def _extract_dynamic_page(self, url: str) -> tuple:
        """Extract content from a dynamic webpage using Selenium."""
        logger.debug(f"Attempting dynamic scrape for URL: {url}")
        try:
            self.driver.get(url)
            self.driver.execute_script("""
                var buttons = document.querySelectorAll("button");
                buttons.forEach(btn => {
                    if (btn.innerText.includes("अधिक पाहा") || btn.innerText.includes("अधिक वाचा") ||
                        btn.innerText.includes("आणखी वाचा") || btn.innerText.includes("अन्खी पहा")) {
                        btn.click();
                    }
                });
                document.querySelectorAll("div, p, span").forEach(el => el.style.display = "block");
            """)
            soup = BeautifulSoup(self.driver.page_source, "html.parser")
            if soup.find("h1") and soup.find("h1").text.strip():
                headline = soup.find("h1").text
            elif soup.find("h2") and soup.find("h2").text.strip():
                headline = soup.find("h2").text
            elif soup.find("h2") and soup.find("h2").find("div") and soup.find("h2").find("div").text.strip():
                headline = soup.find("h2").find("div").text
            else:
                headline = "No headline"
            paragraphs = [p.text for p in soup.find_all("p")]
            article = "\n".join(paragraphs)
            logger.debug(f"Dynamic scrape successful for {url}")
            return headline, article
        except WebDriverException as e:
            logger.warning(f"Dynamic scrape failed for {url}: {str(e)}")
            return None, None
        except Exception as e:
            logger.error(f"Unexpected error in dynamic scrape for {url}: {str(e)}\n{traceback.format_exc()}")
            return None, None

    def _check_leader_in_article(self, article: str) -> bool:
        """Check if the selected leader's name or its variations appear in the article."""
        # LEADERS = self.get_leaders_variations(self.i_file_path, self.leader)
        if not article:
            logger.debug("Article is empty, leader check skipped.")
            return False

        leader_names = []
        for lang in ["English", "Marathi"]:
            leader_names.extend(self.LEADERS[self.leader][lang])
        article_lower = article.lower()
        result = any(name.lower() in article_lower for name in leader_names)
        logger.debug(f"Leader name check for {self.leader}: {'Found' if result else 'Not found'}")
        return result

    def scrape(self):
        """Main method to scrape news articles."""
        logger.info("Starting news scraping process.")
        try:
            # Read Excel file
            logger.debug(f"Reading input file: {self.file_path}")
            try:
                df = pd.read_excel(self.file_path)
                self.news_links = df['Links'].dropna().tolist()
                logger.info(f"Loaded {len(self.news_links)} links from {self.file_path}")
            except Exception as e:
                logger.error(f"Failed to read Excel file {self.file_path}: {str(e)}\n{traceback.format_exc()}")
                raise FileAccessError(f"Failed to read input file: {str(e)}")

            # Validate URLs
            self.news_links = [link for link in self.news_links if is_valid_url(link)]
            logger.debug(f"Filtered to {len(self.news_links)} valid URLs")

            # Setup Selenium driver
            self.link_extraction.setup_driver()

            flag = 0
            if not self.LEADERS:
                logger.info("Leader Name Not available in Name Variations, Cosidering it as Issue Specific Report....")
                flag = 1
            else:
                logger.info("Leader available in Name Variations, Cosidering it as Leader Specific Report....")
                flag = 0

            # Scrape articles
            for link in tqdm(self.news_links, desc="Scraping articles", unit="article"):
                logger.info(f"Processing URL: {link}")
                headline, article = self._extract_static_page(link)
                if not headline or not article:
                    logger.debug(f"Static scrape failed, attempting dynamic scrape for {link}")
                    headline, article = self._extract_dynamic_page(link)

                if headline and article:
                    headline = clean_text(headline)
                    article_cleaned = clean_text(article)
                    if headline in ["Access Denied", "A timeout occurred"]:
                        logger.warning(f"Invalid content (Access Denied/Timeout) for {link}")
                        self.failed_links.append(link)
                        continue
                    if "Verifying you are human" in article_cleaned:
                        logger.warning(f"Human verification detected for {link}")
                        self.failed_links.append(link)
                        continue
                    # Check if leader's name is in the article
                    if flag == 0: # for leader specific
                        if self._check_leader_in_article(article_cleaned):
                            self.extracted_data.append({
                                "Link": link,
                                "Headline": headline,
                                "Article": article_cleaned
                            })
                            logger.info(f"Article extracted successfully for {link}")
                        else:
                            logger.info(f"No mention of {self.leader} in {link}")
                            self.failed_links.append(link)
                    else:  # for issue specific
                        self.extracted_data.append({
                                "Link": link,
                                "Headline": headline,
                                "Article": article_cleaned
                            })
                        logger.info(f"Article extracted successfully for {link}")
                else:
                    logger.warning(f"Failed to extract content from {link}")
                    self.failed_links.append(link)

            # Save results
            logger.info("Saving results.")
            self._save_to_word()
            self._save_failed_links()
            logger.info("Scraping process completed successfully.")

        except Exception as e:
            logger.critical(f"Scraping process failed: {str(e)}\n{traceback.format_exc()}")
            raise NewsScraperError(f"Scraping process failed: {str(e)}")
        finally:
            if self.driver:
                try:
                    self.driver.quit()
                    logger.debug("Selenium WebDriver closed.")
                except Exception as e:
                    logger.error(f"Error closing WebDriver: {str(e)}")

    def _save_to_word(self):
        """Save extracted data to a Word document."""
        logger.info(f"Saving extracted data to Word document: {self.docx_output}")
        try:
            doc = Document()
            heading_style = ensure_heading_style(doc, "Heading 3", 3)

            for idx, item in enumerate(self.extracted_data, start=1):
                doc.add_heading(f"Article {idx}", level=3).style = heading_style
                doc.add_paragraph(f"Title: {item['Headline']}")
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("URL: ")
                run = paragraph.add_run(item['Link'])
                run.font.color.rgb = RGBColor(0, 0, 255)
                part = paragraph.part
                r_id = part.relate_to(item['Link'], RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                hyperlink.append(run._r)
                paragraph._p.append(hyperlink)
                doc.add_paragraph(f"Content: {item['Article']}")

            # Ensure output folder exists
            os.makedirs(self.output_folder, exist_ok=True)
            output_path = os.path.join(self.output_folder, self.docx_output)
            doc.save(output_path)
            logger.debug(f"Word document saved at: {output_path}")
        except Exception as e:
            logger.error(f"Failed to save Word document: {str(e)}\n{traceback.format_exc()}")
            raise FileAccessError(f"Failed to save Word document: {str(e)}")

    def _save_failed_links(self):
        """Save failed links to an Excel file."""
        if self.failed_links:
            logger.info(f"Saving {len(self.failed_links)} failed links to Excel: {self.excel_output}")
            try:
                # Ensure output folder exists
                os.makedirs(self.output_folder, exist_ok=True)
                output_path = os.path.join(self.output_folder, self.excel_output)
                pd.DataFrame({"Failed Links": self.failed_links}).to_excel(output_path, index=False)
                logger.debug(f"Failed links saved at: {output_path}")
            except Exception as e:
                logger.error(f"Failed to save failed links Excel: {str(e)}\n{traceback.format_exc()}")
                raise FileAccessError(f"Failed to save failed links Excel: {str(e)}")
        else:
            logger.debug("No failed links to save.")

    def get_output_paths(self) -> tuple:
        """Return paths to output files."""
        docx_path = os.path.join(self.output_folder, self.docx_output)
        excel_path = os.path.join(self.output_folder, self.excel_output) if self.failed_links else None
        logger.debug(f"Output paths: Word={docx_path}, Excel={excel_path}")
        return docx_path, excel_path